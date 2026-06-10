import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Import settings and repository function
from DomainLinksBackend.app.config import get_settings
from DomainLinksBackend.app.repositories import list_controls_for_branch


def generate_policy_v2():
    # Load settings
    settings = get_settings()
    # Fetch controls for performance-management domain
    controls = list_controls_for_branch(settings, "performance-management")

    # Create a new document
    doc = Document()
    # Set default font
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(11)

    # Header
    header = doc.add_heading('Performance Management Policy', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    metadata_table = doc.add_table(rows=11, cols=2)
    metadata_table.style = 'Table Grid'
    metadata_data = [
        ('Policy Number:', 'SOD-POL-004'),
        ('Version Number:', 'V: 2.00'),
        ('Supersedes:', 'V: 1.00'),
        ('Policy Category:', 'Administrative'),
        ('Approved Date:', '2026-06-01'),
        ('Policy Owner:', 'Executive Strategic Officer, Human Resource Management'),
        ('Governing Body:', 'Finance, Administration and Operations (FAO) Committee'),
        ('Effective Date:', '2026-06-01'),
        ('Policy Author:', 'Human Resource Management Team'),
        ('Review Cycle:', '12 months'),
        ('Inquiries:', 'Executive Strategic Officer, Human Resource Management'),
    ]
    for i, (label, value) in enumerate(metadata_data):
        row = metadata_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Add a simple section introducing controls
    doc.add_heading('2.0 Controls', level=1)
    doc.add_paragraph('The following controls are associated with the Performance Management domain as retrieved from the database.', style='Normal')

    # Create a table of controls
    if controls:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Control Code'
        hdr_cells[1].text = 'Display Name'
        hdr_cells[2].text = 'Description'
        hdr_cells[3].text = 'Objective'
        for ctrl in controls:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ctrl.get('ControlCode', ''))
            row_cells[1].text = str(ctrl.get('DisplayName', ''))
            row_cells[2].text = str(ctrl.get('Description', ''))
            row_cells[3].text = str(ctrl.get('ControlObjective', ''))
    else:
        doc.add_paragraph('No controls found for this domain.', style='Normal')

    # Footer
    doc.add_page_break()
    footer_section = doc.sections[-1]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph('')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('Performance Management Policy\nVersion V: 2.00 | 2026-06-01\n\nFor inquiries, contact: Executive Strategic Officer, Human Resource Management\nReview Cycle: 12 months')
    footer_para.paragraph_format.space_before = Pt(24)

    # Save document
    output_path = os.path.join('WebEnterpriseArc', 'Policies', 'performance-management-policy-v2.00.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'Document saved to {output_path}')

if __name__ == '__main__':
    generate_policy_v2()

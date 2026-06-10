from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create a new document
doc = Document()

# Set default font
doc.styles['Normal'].font.name = 'Segoe UI'
doc.styles['Normal'].font.size = Pt(11)

# Header Section
header = doc.add_heading('Performance Management Policy', 0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add policy metadata table
metadata_table = doc.add_table(rows=11, cols=2)
metadata_table.style = 'Table Grid'

metadata_data = [
    ('Policy Number:', 'SOD-POL-004'),
    ('Version Number:', 'V: 1.00'),
    ('Supersedes:', 'N/A'),
    ('Policy Category:', 'Administrative'),
    ('Approved Date:', '2026-06-01'),
    ('Policy Owner:', 'Executive Strategic Officer, Human Resource Management'),
    ('Governing Body:', 'Finance, Administration and Operations (FAO) Committee'),
    ('Effective Date:', '2026-06-01'),
    ('Policy Author:', 'Human Resource Management Team'),
    ('Review Cycle:', '12 months'),
    ('Inquiries:', 'Executive Strategic Officer, Human Resource Management'),
]

for i, (label, value) in enumerate(metadata_data):
    row = metadata_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()  # Empty paragraph for spacing

# Section 1.0 Context & Objective
doc.add_heading('1.0 Context & Objective', level=1)

doc.add_paragraph(
    '[1.1] Provide the organization with a performance management process that ensures effective governance, '
    'management and continuity for human resource development and employee performance evaluation.',
    style='Normal'
)

doc.add_paragraph(
    '[1.2] Provide a guide on development of performance metrics, weighting methodologies, and management of '
    'employee development plans to support organizational strategic objectives.',
    style='Normal'
)

# Principles
doc.add_heading('Principles', level=2)

principles = [
    '[1.3] Governed and coordinated performance management supports the Human Resource Management domain '
    'strategic objectives for accountability and optimal use of human resources.',
    '[1.4] Performance management begins with the end in mind, focus on scope of work and on outcomes '
    'through formal documentation.',
    '[1.5] Performance initiatives have a business purpose that provides value to the organization through '
    'employee development and performance optimization.',
    '[1.6] Performance management is planned, scheduled, tracked, and documented and includes cost estimation '
    'and performance parameters for all employees.',
    '[1.7] Investment in the use of human resources carries the expectation of future return for operational '
    'development, infrastructure, assets or the acquisition of knowledge or capacity through structured '
    'development plans.',
]

for principle in principles:
    doc.add_paragraph(principle, style='Normal')

# Accountability and Transparency
doc.add_heading('Accountability and Transparency', level=2)

accountability = [
    '[1.8] Performance documentation will be made accessible by the Human Resource Management team to the '
    'employees, supervisors and leadership.',
    '[1.9] Performance documentation accessibility support accountability and improves transparency for human '
    'resources used and for employee work at any given time.',
    '[1.10] Governance, measurement, and accountability are disciplines imposed by performance management processes.',
    '[1.11] A Performance Management Framework describes the governance and oversight structure for a '
    'performance management investment including time, human and financial resources.',
]

for item in accountability:
    doc.add_paragraph(item, style='Normal')

# Strategy
doc.add_heading('Strategy', level=2)

strategy = [
    '[1.12] The Human Resource Management domain identified an organizational responsibility to apply '
    'performance-based planning and establish clear timelines for employee development.',
    '[1.13] The Human Resource Management domain identified that performance management requests must be '
    'reviewed and approved through an established process.',
    '[1.14] Performance Management is a pillar of the organizational governance framework. The objective is to '
    'partner for the outcome of collaboration and transparency in employee development.',
    '[1.15] Performance Management is a pillar of the Human Resource Management domain. The objective is to '
    'develop and implement the foundational frameworks that support the core components of the organization, '
    'of which performance management is one framework.',
]

for item in strategy:
    doc.add_paragraph(item, style='Normal')

# Section 2.0 Application
doc.add_heading('2.0 Application', level=1)

doc.add_paragraph(
    '[2.1] This policy applies to all Human Resource Management employees who are responsible in any capacity '
    'for development, approval and implementation of performance management initiatives.',
    style='Normal'
)

# Section 3.0 Policy
doc.add_heading('3.0 Policy', level=1)

doc.add_paragraph(
    '[3.1] All Human Resource Management performance initiatives shall be managed through this policy.',
    style='Normal'
)

doc.add_paragraph(
    '[3.2] Human Resource Management Key Leadership has the responsibility and authority to enforce this policy.',
    style='Normal'
)

# Mandatory Signed Role Approval
doc.add_heading('Mandatory Signed Role Approval', level=2)

role_approval = [
    '[3.3] A Mandatory Signed Role Approval must be used as an option to a Performance Development Plan or '
    'Performance Framework.',
    '[3.4] A performance proposal must begin with documented role requirements as directed by the unit manager '
    'or supervisor prior to development of employee performance metrics.',
    '[3.5] A Performance Management Framework must be used to document employee performance, time and scope, '
    'identify business outcomes or deliverables to ensure that the organization is delivered what it is seeking '
    'to achieve through employee development.',
]

for item in role_approval:
    doc.add_paragraph(item, style='Normal')

# Annual Weighting Calibration Review
doc.add_heading('Annual Weighting Calibration Review', level=2)

weighting_calibration = [
    '[3.6] Managers must review any performance submissions for approval to ensure the performance management '
    'represents an effective and efficient solution to the business need or opportunity and is one that supports '
    "the organization's strategic direction.",
    '[3.6.1] A requirement must be included in performance documentation by a manager that will determine what '
    'level of financial resources the performance initiative needs, and will identify the finance source (unit '
    'budget/Finance, Administration and Operations Committee).',
    '[3.7] A Performance Close-Out Report must be written as part of performance management close out.',
]

for item in weighting_calibration:
    doc.add_paragraph(item, style='Normal')

# Development Plan Weighting Methodology Definition
doc.add_heading('Development Plan Weighting Methodology Definition', level=2)

methodology = [
    '[3.8] A Development Plan Weighting Methodology shall be used to describe an assignment purpose, milestones '
    'and deliverables for employee development.',
    '[3.9] The unit manager or supervisor shall determine if a Development Plan Weighting Methodology is used.',
    '[3.10] A Development Plan Weighting Methodology must use established methodology to compliment or support '
    'an existing performance initiative or to document a general assignment.',
    '[3.11] The unit manager or supervisor shall document and track Development Plan Weighting Methodologies '
    'using the established methodology framework.',
]

for item in methodology:
    doc.add_paragraph(item, style='Normal')

# Role Weighting Deviation Justification Record
doc.add_heading('Role Weighting Deviation Justification Record', level=2)

deviation = [
    '[3.12] If a Performance Framework will be used to document a performance initiative, then a Role Weighting '
    'Deviation Justification Record shall be produced.',
    '[3.13] The unit manager or supervisor shall determine use of a Role Weighting Deviation Justification '
    'Record to describe a proposal for a performance initiative.',
    '[3.14] A Role Weighting Deviation Justification Record must be used to describe a proposal for a '
    'performance initiative that identifies a business need, performance purpose, options analysis and risk '
    'assessment as justification to produce deliverables.',
]

for item in deviation:
    doc.add_paragraph(item, style='Normal')

# Development Plan Weighting Methodology details
doc.add_paragraph(
    'A Development Plan Weighting Methodology:', style='Normal'
)

methodology_details = [
    'can support a task linked to an existing performance initiative;',
    'is used for a performance initiative when the scope is limited;',
    'is used for a performance initiative when the performance initiative will use minimal resources to achieve its goal;',
    'documents any task assigned when not related to a performance initiative;',
    'is work to be done, and is documented for purpose, milestones, deliverable and is signed by a manager or supervisor.',
]

for detail in methodology_details:
    p = doc.add_paragraph()
    p.add_run(detail)
    p.paragraph_format.left_indent = Inches(0.5)

# Manager Alignment Certification on Development Weightings
doc.add_heading('Manager Alignment Certification on Development Weightings', level=2)

certification = [
    '[3.15] Performance outcomes shall be clearly defined, measurable and developed with stakeholder involvement.',
    '[3.16] A Development Plan Weighting Methodology shall be approved by the developer\'s manager or unit supervisor.',
    '[3.16.1] The unit budget may have the resources necessary to carry out the performance initiative through '
    'budget reallocation; this is the first line of research for financing the performance initiative.',
    '[3.16.2] Dependent on the performance initiative\'s financial need when it is determined that the unit '
    'budget does not have the resource for reallocation, the Development Plan Weighting Methodology shall be '
    'presented to the Finance, Administration and Operations Committee to request funding support.',
    '[3.16.3] The performance initiative financial requirements must be secured prior to final decision to '
    'proceed with a Performance Framework.',
    '[3.17] Approval of a Development Plan Weighting Methodology must occur prior to a Performance Framework.',
]

for item in certification:
    doc.add_paragraph(item, style='Normal')

# Mandatory Three-Year Employee Development Plan
doc.add_heading('Mandatory Three-Year Employee Development Plan', level=2)

three_year_plan = [
    '[3.18] A Mandatory Three-Year Employee Development Plan must be completed using established methodology '
    'for all performance initiatives.',
    '[3.19] An approved Development Plan Weighting Methodology must precede development of a Mandatory '
    'Three-Year Employee Development Plan.',
    '[3.20] A Mandatory Three-Year Employee Development Plan shall be developed in collaboration with a '
    'performance manager and a performance sponsor.',
]

for item in three_year_plan:
    doc.add_paragraph(item, style='Normal')

# Approval Process
doc.add_heading('Approval Process', level=2)

approval_process = [
    '[3.21] An approval body or person must be identified for approval of a Mandatory Three-Year Employee '
    'Development Plan.',
    '[3.22] Additional approval requirements may apply based on the scope and complexity of the performance initiative.',
    '[3.23] A Performance Manager shall obtain unit senior management support for the performance initiative.',
]

for item in approval_process:
    doc.add_paragraph(item, style='Normal')

# Performance Communications
doc.add_heading('Performance Communications', level=2)

communications = [
    '[3.24] The Performance Manager shall communicate performance information including progress milestones to '
    'unit senior management and identified stakeholders.',
    '[3.25] A Close-Out Report (COR) must be written by the performance manager when the performance initiative is completed.',
    '[3.26] A COR must be distributed to the Human Resource Management designate and to those identified by '
    'the supervisor to receive copy.',
    '[3.27] Performance management documentation must be archived following the Records Management Department established policy.',
]

for item in communications:
    doc.add_paragraph(item, style='Normal')

# Quality Assurance
doc.add_heading('Quality Assurance', level=2)

quality_assurance = [
    '[3.28] The Human Resource Management domain shall track performance documents for quality assurance.',
    '[3.29] All drafted Performance Frameworks must be given to a Human Resource Management designate for '
    'quality review and tracking prior to its approval.',
    '[3.30] The Performance Manager shall submit the approved Performance Framework to Human Resource '
    'Management for tracking and monitoring of performance milestone progress.',
    '[3.31] The Human Resource Management designate shall ensure that the senior manager/Portfolio Chief are '
    'informed of performance progress, that includes the performance framework, milestones, performance '
    'completion and a performance close out report.',
]

for item in quality_assurance:
    doc.add_paragraph(item, style='Normal')

# Section 4.0 Consequences
doc.add_heading('4.0 Consequences', level=1)

consequences = [
    '[4.1] Human Resource Management will address non-compliance to this policy by implementing corrective '
    'measures. Notification to Unit key leadership will be sent when there is non-compliance to this policy.',
    '[4.2] Performance Frameworks shall not move forward in the approval process if not in compliance with standards.',
    '[4.3] Non-compliance to the Performance Management Policy may have disciplinary measures applied as '
    'outlined in the Human Resource Management Administration-Personnel Policy and the Performance Management '
    'Policy from Human Resources.',
]

for item in consequences:
    doc.add_paragraph(item, style='Normal')

# Section 5.0 Foundational Records & Documents
doc.add_heading('5.0 Foundational Records & Documents', level=1)

foundational_records = [
    '[5.1] Human Resource Management Domain Framework (2026) - As a result of domain framework development '
    'sessions, a summary was developed. This report identified values and guiding principles that stated to '
    'foster collaboration and be stakeholder inclusive and value stability through minimum standards. Key '
    'organizational objectives identified for operations were to develop systems and mechanisms and to develop '
    'measurement and tracking systems.',
    '[5.2] Human Resource Management Domain Design Framework Final Report (2026) - As a result of domain '
    'framework development, an organizational design framework was established and through consultation with '
    'specific groups, limits that are non-negotiable were identified. A limit, Process/Systems, stated that '
    '"performance requests must be reviewed and approved through an established process". This policy supports '
    'the limits set out by the Domain Framework Development.',
    '[5.3] Human Resource Management and Performance Development are mandate domains that support this policy.',
]

for item in foundational_records:
    doc.add_paragraph(item, style='Normal')

# Section 6.0 Definitions
doc.add_heading('6.0 Definitions', level=1)

doc.add_heading('Abbreviations & Acronyms', level=2)

abbreviations = [
    '[6.1] [COR]: Close-Out Report - A documented report of performance success and failures, lessons learned, '
    'and one that formalizes the completion of a performance initiative.',
    '[6.2] [FAO]: Finance, Administration and Operations Committee - The governing body responsible for '
    'reviewing and approving performance initiatives requiring funding support.',
    '[6.3] [HRIS]: Human Resource Information System - The system repository where performance documentation '
    'and approval records are maintained.',
]

for item in abbreviations:
    doc.add_paragraph(item, style='Normal')

doc.add_heading('Key Definitions', level=2)

definitions = [
    '[6.4] [Development Plan Weighting Methodology]: A formal framework requiring specific development plan '
    'tiers to correspond with defined performance weighting ranges, ensuring consistent application of '
    'weightings across roles based on strategic development goals.',
    '[6.5] [Mandatory Signed Role Approval]: A requirement that all employees must have a signed Role & '
    'Responsibility (R&R) document on file before performance metrics are set, ensuring performance '
    'expectations are formally documented and agreed upon prior to evaluation periods.',
    '[6.6] [Annual Weighting Calibration Review]: An annual review process of weighting logic to ensure '
    'alignment with current organizational strategy, maintaining relevance and accuracy of performance '
    'weights as business priorities shift.',
    '[6.7] [Role Weighting Deviation Justification Record]: A documented rationale required when weighting '
    'allocation deviates from standard matrices due to development plan changes, preventing arbitrary '
    'weighting adjustments by requiring traceable links between development needs and performance expectations.',
    '[6.8] [Manager Alignment Certification on Development Weightings]: An annual certification requirement '
    'for managers to confirm that assigned role weightings align with approved individual development plans, '
    'promoting accountability and transparency in how career development strategies influence performance '
    'evaluation criteria.',
    '[6.9] [Mandatory Three-Year Employee Development Plan]: A formal three-year development roadmap required '
    'for all eligible employees, ensuring long-term career planning is initiated and documented across the workforce.',
    '[6.10] [Strategic Alignment Verification on Development Milestones]: A requirement that development '
    'milestones be explicitly mapped against organizational strategic objectives, aligning individual growth '
    'trajectories with corporate strategic direction.',
]

for item in definitions:
    doc.add_paragraph(item, style='Normal')

# Section 7.0 Forms & Templates
doc.add_heading('7.0 Forms & Templates', level=1)

forms_templates = [
    '[7.1] [Role & Responsibility Approval Template] (ANNEX A) - Template for documenting signed role '
    'approvals and performance expectations.',
    '[7.2] [Development Plan Weighting Methodology Document] (ANNEX B) - Framework document outlining the '
    'mapping logic between development plan levels and role weighting percentages.',
    '[7.3] [Three-Year Employee Development Plan Template] (ANNEX C) - Comprehensive template for creating '
    'formal three-year development roadmaps for employees.',
]

for item in forms_templates:
    doc.add_paragraph(item, style='Normal')

# Section 8.0 References
doc.add_heading('8.0 References', level=1)

references = [
    '[8.1] [Guideline: Using the Performance Management Framework for Your Initiative] - Comprehensive '
    'guidelines for implementing performance management processes.',
    '[8.2] [Human Resource Management Domain Framework (2026)] - Foundational document establishing the '
    'domain structure and governance requirements.',
    '[8.3] [Performance Management Policy Guidelines] - Detailed guidelines for policy implementation and compliance.',
]

for item in references:
    doc.add_paragraph(item, style='Normal')

# Footer
doc.add_page_break()
footer_section = doc.sections[-1]
footer = footer_section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph('')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run('Performance Management Policy\nVersion V: 1.00 | 2026-06-01\n\nFor inquiries, contact: Executive Strategic Officer, Human Resource Management\nReview Cycle: 12 months')
footer_para.paragraph_format.space_before = Pt(24)

# Save the document
output_file = r'c:\Users\richardbasque\repo\DomainLinksRepo\WebEnterpriseArc\Policies\performance-management-policy-v1.00.docx'
doc.save(output_file)

print(f"Document saved successfully to: {output_file}")

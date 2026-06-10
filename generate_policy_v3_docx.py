import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Import settings and repository function to fetch controls
from DomainLinksBackend.app.config import get_settings
from DomainLinksBackend.app.repositories import list_controls_for_branch

def generate_policy_v3():
    settings = get_settings()
    # Fetch controls for the performance-management domain
    controls = list_controls_for_branch(settings, "performance-management")

    doc = Document()
    # Set default font
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(11)

    # Header
    header = doc.add_heading('Performance Management Policy', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table (11 rows)
    meta = [
        ('Policy Number:', 'SOD-POL-004'),
        ('Version Number:', 'V: 3.00'),
        ('Supersedes:', 'V: 2.00'),
        ('Policy Category:', 'Administrative'),
        ('Approved Date:', datetime.now().strftime('%Y-%m-%d')),
        ('Policy Owner:', 'Executive Strategic Officer, Human Resource Management'),
        ('Governing Body:', 'Finance, Administration and Operations (FAO) Committee'),
        ('Effective Date:', datetime.now().strftime('%Y-%m-%d')),
        ('Policy Author:', 'Human Resource Management Team'),
        ('Review Cycle:', '12 months'),
        ('Inquiries:', 'Executive Strategic Officer, Human Resource Management'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metadata'
    hdr_cells[1].text = 'Value'
    for label, value in meta:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Section 1.0 Context & Objective
    doc.add_heading('1.0 Context & Objective', level=1)
    doc.add_paragraph('Provide the organization with a performance management process that ensures effective governance, management and continuity for employee performance evaluation and development.', style='Normal')
    doc.add_paragraph('Guide the development of performance metrics, weighting methodologies, and employee development plans to support strategic objectives.', style='Normal')

    # Principles
    doc.add_heading('Principles', level=2)
    principles = [
        'Governed and coordinated performance management supports strategic objectives for accountability and optimal use of human resources.',
        'Performance initiatives begin with the end in mind, focusing on scope of work and outcomes.',
        'Performance initiatives have a business purpose that provides value to the organization.',
        'Performance initiatives are planned, scheduled, tracked, and documented, including cost estimation and performance parameters.',
        'Investment in performance management carries the expectation of future return for operational development and organizational capability.',
    ]
    for p in principles:
        doc.add_paragraph(p, style='List Bullet')

    # Accountability and Transparency
    doc.add_heading('Accountability and Transparency', level=2)
    accountability = [
        'Performance documentation will be made accessible by Human Resource Management to employees, supervisors and leadership.',
        'Accessibility supports accountability and improves transparency for resources used and work performed.',
        'Governance, measurement, and accountability are disciplines imposed by the performance management process.',
        'A policy document describes the governance and oversight structure for performance management investments including time, human and financial resources.',
    ]
    for a in accountability:
        doc.add_paragraph(a, style='List Bullet')

    # Strategy
    doc.add_heading('Strategy', level=2)
    strategy = [
        'Human Resource Management identified an organizational responsibility to apply performance‑based planning and establish clear timelines.',
        'Performance‑management requests must be reviewed and approved through an established process.',
        'Performance Management is a pillar of the organizational governance framework. The objective is to partner for collaboration and transparency.',
        'Performance Management supports the foundational frameworks that enable the organization to achieve its strategic goals.',
    ]
    for s in strategy:
        doc.add_paragraph(s, style='List Bullet')

    # Section 2.0 Application
    doc.add_heading('2.0 Application', level=1)
    doc.add_paragraph('This policy applies to all Human Resource Management employees who are responsible for development, approval and implementation of performance management initiatives.', style='Normal')

    # Section 3.0 Policy
    doc.add_heading('3.0 Policy', level=1)
    doc.add_paragraph('All performance‑management initiatives shall be managed through this policy.', style='Normal')
    doc.add_paragraph('Human Resource Management Key Leadership has the responsibility and authority to enforce this policy.', style='Normal')

    # Insert policy statements for each control
    doc.add_heading('Performance Management Controls', level=2)
    if controls:
        for ctrl in controls:
            display_name = str(ctrl.get('DisplayName', ''))
            description = str(ctrl.get('Description', '')).strip()
            objective = str(ctrl.get('ControlObjective', '')).strip()
            evidence = str(ctrl.get('EvidenceExpectation', '')).strip()
            # Control heading
            doc.add_heading(display_name, level=3)
            # Statement 1: Requirement
            stmt1 = f"The organization shall implement the control '{display_name}' as defined."
            if description:
                stmt1 += f" {description}"
            doc.add_paragraph(stmt1, style='Normal')
            # Statement 2: Objective (if available)
            if objective:
                stmt2 = f"The purpose of this control is to {objective.lower()}"
                doc.add_paragraph(stmt2, style='Normal')
            # Statement 3: Evidence (if available)
            if evidence:
                stmt3 = f"Compliance shall be demonstrated by {evidence.lower()}."
                doc.add_paragraph(stmt3, style='Normal')
    else:
        doc.add_paragraph('No controls were found for the Performance Management domain.', style='Normal')

    # Additional generic policy items (retain from v2 for completeness)
    doc.add_heading('Mandatory Signed Role Approval', level=2)
    doc.add_paragraph('All employees must have a signed Role & Responsibility (R&R) document on file before performance metrics are set.', style='Normal')
    doc.add_heading('Annual Weighting Calibration Review', level=2)
    doc.add_paragraph('An annual review of weighting logic must be conducted to ensure alignment with current organizational strategy.', style='Normal')
    doc.add_heading('Development Plan Weighting Methodology Definition', level=2)
    doc.add_paragraph('Formal rules require specific development plan tiers to correspond with defined performance weighting ranges.', style='Normal')
    doc.add_heading('Role Weighting Deviation Justification Record', level=2)
    doc.add_paragraph('Documented rationale is required when weighting allocation deviates from standard matrices due to development plan changes.', style='Normal')
    doc.add_heading('Manager Alignment Certification on Development Weightings', level=2)
    doc.add_paragraph('Managers must certify annually that assigned role weightings align with approved individual development plans.', style='Normal')
    doc.add_heading('Mandatory Three‑Year Employee Development Plan', level=2)
    doc.add_paragraph('A formal three‑year development roadmap must be created for all eligible employees.', style='Normal')
    doc.add_heading('Strategic Alignment Verification on Development Milestones', level=2)
    doc.add_paragraph('Development milestones must be explicitly mapped against organizational strategic objectives.', style='Normal')
    doc.add_heading('Annual Development Plan Refresh Cycle', level=2)
    doc.add_paragraph('Annual review and updating of the three‑year development plan is required to reflect current career status.', style='Normal')

    # Section 4.0 Consequences
    doc.add_heading('4.0 Consequences', level=1)
    doc.add_paragraph('Human Resource Management will address non‑compliance by implementing corrective measures and notifying unit leadership.', style='Normal')
    doc.add_paragraph('Non‑compliant initiatives shall not progress in the approval process.', style='Normal')
    doc.add_paragraph('Disciplinary measures may be applied as outlined in the HR Administration‑Personnel Policy.', style='Normal')

    # Section 5.0 Foundational Records & Documents
    doc.add_heading('5.0 Foundational Records & Documents', level=1)
    doc.add_paragraph('Performance Management Policy (v2.00) – Provides the baseline governance framework.', style='Normal')
    doc.add_paragraph('Strategic Plan 2026 – Aligns performance initiatives with organizational objectives.', style='Normal')
    doc.add_paragraph('HR Handbook – Defines roles, responsibilities and documentation requirements.', style='Normal')

    # Section 6.0 Definitions
    doc.add_heading('6.0 Definitions', level=1)
    definitions = {
        'Role & Responsibility (R&R) Document': 'A signed record outlining employee duties and performance expectations.',
        'Weighting Calibration': 'The process of reviewing and adjusting performance weighting to reflect strategic priorities.',
        'Development Plan': 'A structured roadmap for employee skill growth and career progression over a defined period.',
        'Control': 'A specific requirement or guideline that governs performance management activities.',
    }
    for term, definition in definitions.items():
        doc.add_paragraph(f'{term}: {definition}', style='Normal')

    # Section 7.0 Forms & Templates
    doc.add_heading('7.0 Forms & Templates', level=1)
    forms = [
        'Signed Role & Responsibility (R&R) Template (ANNEX A)',
        'Annual Weighting Calibration Agenda (ANNEX B)',
        'Three‑Year Development Plan Template (ANNEX C)',
    ]
    for f in forms:
        doc.add_paragraph(f, style='List Bullet')

    # Section 8.0 References
    doc.add_heading('8.0 References', level=1)
    refs = [
        'Strategic Plan 2026 – Organizational strategic objectives.',
        'HR Handbook – Policies and procedures for employee management.',
        'Performance Management Guidelines – Detailed implementation guidance.',
    ]
    for r in refs:
        doc.add_paragraph(r, style='List Bullet')

    # Footer
    doc.add_page_break()
    footer_section = doc.sections[-1]
    footer = footer_section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph('')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Performance Management Policy\nVersion V: 3.00 | ' + datetime.now().strftime('%Y-%m-%d') + '\n\nFor inquiries, contact: Executive Strategic Officer, Human Resource Management\nReview Cycle: 12 months')
    para.paragraph_format.space_before = Pt(24)

    output_path = os.path.join('WebEnterpriseArc', 'Policies', 'performance-management-policy-v3.00.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'Document saved to {output_path}')

if __name__ == '__main__':
    generate_policy_v3()
